import streamlit as st
import pandas as pd
import io
import copy
from datetime import datetime
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.oxml.ns import qn

# ==========================================
# 页面配置（增强）
# ==========================================
st.set_page_config(
    page_title="PRER报告生成系统",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# CSS（增强版 UI）
# ==========================================
st.markdown("""
<style>
.main { background-color: #f8f9fb; }
h1, h2, h3 { color: #2c3e50; }

.stButton>button {
    width: 100%;
    border-radius: 12px;
    height: 3em;
    font-weight: 600;
}

.stDownloadButton>button {
    width: 100%;
    border-radius: 12px;
    background-color: #27ae60;
    color: white;
    font-weight: 600;
}

.block-container {
    padding-top: 2rem;
}

[data-testid="stExpander"] {
    border-radius: 10px;
    border: 1px solid #e6e6e6;
}
</style>
""", unsafe_allow_html=True)

# ==========================================
# 核心函数（完全保留）
# ==========================================
def format_date(value):
    if pd.isna(value): return ""
    if isinstance(value, (pd.Timestamp, datetime)):
        return f"{value.year}年{value.month}月{value.day}日"
    return str(value)

def set_dual_font_10pt(run):
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')

def replace_text_strict_format(paragraph, replacements):
    if not paragraph.runs: return
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for k, v in replacements.items():
        new_text = new_text.replace(str(k), str(v))
    if new_text != full_text:
        paragraph.runs[0].text = new_text
        for i in range(1, len(paragraph.runs)):
            paragraph.runs[i].text = ""

def replace_all_content_keep_style(doc, replacements):
    for p in doc.paragraphs:
        replace_text_strict_format(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_strict_format(p, replacements)
    for node in doc.element.xpath("//w:t"):
        if node.text:
            text = node.text
            for k, v in replacements.items():
                text = text.replace(str(k), str(v))
            node.text = text

def replace_header_keep_format(doc, replacements):
    def replace_in_header(header):
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if not run.text: continue
                text = run.text
                for k, v in replacements.items():
                    if str(k) in text:
                        text = text.replace(str(k), str(v))
                run.text = text

    for section in doc.sections:
        try:
            section.header.is_linked_to_previous = False
            replace_in_header(section.header)
            if hasattr(section, 'first_page_header'):
                replace_in_header(section.first_page_header)
            if hasattr(section, 'even_page_header'):
                replace_in_header(section.even_page_header)
        except:
            pass

def fill_specific_table(table, records):
    if not records or len(table.rows) < 2: return
    template_row = table.rows[1]

    while len(table.rows) > 2:
        table._tbl.remove(table.rows[-1]._tr)

    for idx, record in enumerate(records, start=1):
        new_row_xml = copy.deepcopy(template_row._tr)
        table._tbl.append(new_row_xml)
        row = table.rows[-1]

        row_data = [
            str(idx),
            record[1],
            record[2],
            record[3],
            record[4]
        ]

        for i, val in enumerate(row_data):
            if i < len(row.cells):
                cell = row.cells[i]
                if not cell.paragraphs:
                    cell.add_paragraph()

                p = cell.paragraphs[0]
                p.text = ""

                run = p.add_run(str(val) if not pd.isna(val) else "")
                set_dual_font_10pt(run)

                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table._tbl.remove(template_row._tr)

# ==========================================
# Sidebar（增强）
# ==========================================
with st.sidebar:
    st.title("⚙️ 参数设置")

    st.markdown("### 📊 表格索引配置")
    st.caption("⚠️ 从0开始")

    idx_wanfang = st.number_input("万方", value=3)
    idx_cnki = st.number_input("知网", value=4)
    idx_pubmed = st.number_input("PubMed", value=5)
    idx_embase = st.number_input("Embase", value=6)

    st.divider()

    st.markdown("### ℹ️ 使用说明")
    st.info("""
必须包含：
- Sheet1：placeholder / value
- Sheet2：数据库 / 文献信息
""")

# ==========================================
# 主界面
# ==========================================
st.title("📄 PRER报告自动生成系统")
st.caption("CER中心自动化工具｜3秒生成Word报告")

col1, col2 = st.columns(2)

with col1:
    template_file = st.file_uploader("上传Word模板", type=["docx"])

with col2:
    excel_file = st.file_uploader("上传Excel数据", type=["xlsx"])

# ==========================================
# Excel预览（增强）
# ==========================================
if excel_file:
    try:
        xl = pd.ExcelFile(excel_file)
        with st.expander("🔍 Excel预览"):
            sheet = st.selectbox("选择Sheet", xl.sheet_names)
            df_preview = pd.read_excel(excel_file, sheet_name=sheet)
            st.dataframe(df_preview.head(10), use_container_width=True)
    except Exception as e:
        st.error(f"Excel解析失败: {e}")

st.divider()

# ==========================================
# 核心执行（增强稳定性）
# ==========================================
if st.button("🚀 生成报告"):

    if not template_file or not excel_file:
        st.warning("请上传完整文件")
        st.stop()

    try:
        with st.spinner("处理中..."):

            # 读取Excel
            xl = pd.ExcelFile(excel_file)

            if len(xl.sheet_names) < 2:
                st.error("Excel至少需要2个Sheet")
                st.stop()

            p_df = pd.read_excel(excel_file, sheet_name=xl.sheet_names[0])
            l_df = pd.read_excel(excel_file, sheet_name=xl.sheet_names[1])

            # 占位符
            replacements = {}
            for _, row in p_df.iterrows():
                key = str(row.get("placeholder", "")).strip()
                if not key: continue

                if not key.startswith("{"): key = "{" + key
                if not key.endswith("}"): key = key + "}"

                replacements[key] = format_date(row.get("value", ""))

            # 文献分类
            databases = {}
            for _, row in l_df.iterrows():
                db_raw = str(row.get("数据库", "")).upper()

                if "万方" in db_raw:
                    db = "万方"
                elif "知网" in db_raw or "CNKI" in db_raw:
                    db = "中国知网"
                elif "PUBMED" in db_raw:
                    db = "Pubmed"
                elif "EMBASE" in db_raw:
                    db = "Embase"
                else:
                    continue

                record = [
                    str(row.get("文献编号", "")),
                    str(row.get("文献信息", "")),
                    str(row.get("检索说明", "")),
                    str(row.get("排除理由", "")),
                    str(row.get("备注", ""))
                ]

                databases.setdefault(db, []).append(record)

            total = sum(len(v) for v in databases.values())
            replacements["{文献量}"] = str(total)
            replacements["{总纳入文献量}"] = str(total)

            # Word处理
            doc = Document(template_file)

            replace_all_content_keep_style(doc, replacements)
            replace_header_keep_format(doc, replacements)

            # 表格填充
            db_map = {
                idx_wanfang: "万方",
                idx_cnki: "中国知网",
                idx_pubmed: "Pubmed",
                idx_embase: "Embase"
            }

            for idx, db_name in db_map.items():
                if idx < len(doc.tables):
                    fill_specific_table(doc.tables[idx], databases.get(db_name, []))

            # 输出
            out = io.BytesIO()
            doc.save(out)
            out.seek(0)

        st.success("✅ 生成成功")
        st.balloons()

        st.download_button(
            "📥 下载报告",
            data=out,
            file_name=f"PRER报告_{datetime.now().strftime('%m%d_%H%M')}.docx"
        )

    except Exception as e:
        st.error(f"生成失败: {e}")

# --- 版权信息 ---
st.divider()
st.caption("© CER中心 自动化排版工作室 | Seven")
