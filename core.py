# -*- coding: utf-8 -*-

import os
import re
import copy
import pandas as pd
from datetime import datetime
from docx import Document

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn


# =========================
# 数据读取（五表结构升级版｜文件级输入）
# =========================

def load_table(file):
    if file is None:
        return None
    return pd.read_excel(file)


def load_all_tables(file_map):
    """
    统一加载五个标准数据表（文件级）

    最低要求：
        - 产品信息表（必须）
        - 文献筛选表（必须）
    """

    data = {
        "product_info": load_table(file_map.get("product_info")),
        "screening": load_table(file_map.get("screening")),
        "analysis": load_table(file_map.get("analysis")),
        "efficacy": load_table(file_map.get("efficacy")),
        "safety": load_table(file_map.get("safety")),
    }

    # =========================
    # ⭐ 强制校验（核心逻辑）
    # =========================

    if data["product_info"] is None or data["product_info"].empty:
        raise ValueError("❌ 缺少【产品信息表】或文件为空")

    if data["screening"] is None or data["screening"].empty:
        raise ValueError("❌ 缺少【文献筛选表】或文件为空")

    return data

def format_date(value):
    if pd.isna(value):
        return ""
    if isinstance(value, (pd.Timestamp, datetime)):
        return f"{value.year}年{value.month}月{value.day}日"
    return str(value)


def load_placeholders_from_product_table(product_df):
    """
    从产品信息表读取Word占位符
    """

    if product_df is None or product_df.empty:
        raise ValueError("❌ 产品信息表不存在或为空")

    required_cols = {"placeholder", "value"}

    if not required_cols.issubset(set(product_df.columns)):
        raise ValueError("❌ 产品信息表必须包含列：placeholder, value")

    replacements = {}

    for _, row in product_df.iterrows():
        key = str(row.get("placeholder", "")).strip()
        value = format_date(row.get("value", ""))

        if not key:
            continue

        if not key.startswith("{"):
            key = "{" + key
        if not key.endswith("}"):
            key += "}"

        replacements[key] = value.strip()

    return replacements


# =========================
# ⭐ 文献读取（修复Cochrane）
# =========================
def load_literature_from_screening(screening_df):
    """
    从文献筛选表生成数据库结构
    """

    if screening_df is None:
        return {}

    databases = {}

    for _, row in screening_df.iterrows():
        db_raw = str(row.get("数据库", "")).strip().upper()

        if "万方" in db_raw:
            db = "万方"
        elif "知网" in db_raw or "CNKI" in db_raw:
            db = "中国知网"
        elif "PUBMED" in db_raw:
            db = "PubMed"
        elif "COCHRANE" in db_raw:
            db = "Cochrane Library"
        else:
            continue

        record = [
            str(row.get("文献编号", "")),
            str(row.get("文献信息", "")),
            str(row.get("检索说明", "")),
            str(row.get("排除理由", "")),
            str(row.get("备注", ""))
        ]

        databases.setdefault(db, []).append(record)

    return databases



# =========================
# ⭐ 纳入判断
# =========================
def is_included(text):
    if not text:
        return False
    return "纳入" in str(text)


# ⭐ 生成统计描述（核心升级）
# =========================
def build_summary_text(db_map, product_name):

    included_map = {}

    for db, records in db_map.items():

        ids = []
        new_index = 1

        for r in records:
            if is_included(r[2]):
                ids.append(str(new_index))
            new_index += 1

        if ids:
            included_map[db] = ids

    total = sum(len(v) for v in included_map.values())

    # 无纳入
    if total == 0:
        return 0, "未检索到符合纳入标准的文献。"

    parts = []

    for db, ids in included_map.items():
        parts.append(f"{db}（{'、'.join(ids)}）")

    # 多数据库
    if len(parts) > 1:
        db_text = "和".join(parts)
        prefix = "分别在"
        suffix = "检索获得。"

    # 单数据库
    else:
        db = list(included_map.keys())[0]
        ids = included_map[db]

        prefix = f"在{db}检索获得，文献编号"
        suffix = "。"
        db_text = "、".join(ids)

    text = f"本次检索纳入分析的文献为{total}篇，{prefix}{db_text}{suffix}针对文献中{product_name}（以下简称为“目标产品”）临床使用的安全有效性进行综述分析，文献筛选详见附录1（文献数据分析）。"

    return total, text

def build_literature_analysis_text(analysis_df, screening_df):
    import re

    if analysis_df is None or analysis_df.empty:
        return "未提供文献数据分析内容。"

    if screening_df is None or screening_df.empty:
        return "缺少文献筛选表，无法建立编号映射。"

    # =========================
    # 1️⃣ 构建编号映射
    # =========================
    id_map = {}
    db_groups = {}

    for _, row in screening_df.iterrows():
        db_raw = str(row.get("数据库", "")).strip().upper()

        if "万方" in db_raw:
            db = "万方"
        elif "知网" in db_raw or "CNKI" in db_raw:
            db = "中国知网"
        elif "PUBMED" in db_raw:
            db = "PubMed"
        elif "COCHRANE" in db_raw:
            db = "Cochrane Library"
        else:
            continue

        db_groups.setdefault(db, []).append(row)

    for db, rows in db_groups.items():
        for idx, row in enumerate(rows, start=1):
            old_id = str(row.get("文献编号", "")).strip()
            id_map[old_id] = idx

    # =========================
    # 2️⃣ 替换函数
    # =========================
    def replace_doc_id_in_text(text):
        def repl(match):
            old_id = match.group(1)
            new_id = id_map.get(old_id)
            return f"文献{new_id}" if new_id else match.group(0)

        return re.sub(r"文献\s*(\d+)", repl, text)

    # =========================
    # 3️⃣ 清洗分析表
    # =========================
    df = analysis_df.dropna(subset=["文献编号", "文献数据分析描述"]).copy()
    df["文献编号"] = df["文献编号"].astype(str)

    df["新编号"] = df["文献编号"].map(id_map)
    df = df.dropna(subset=["新编号"])

    if df.empty:
        return "未匹配到有效文献分析内容。"

    df["新编号"] = df["新编号"].astype(int)
    df = df.sort_values(by="新编号")

    # =========================
    # 4️⃣ 分组拼接
    # =========================
    grouped = df.groupby("新编号")["文献数据分析描述"].apply(list)

    parts = []

    for _, texts in grouped.items():
        processed_texts = []

        for t in texts:
            t = str(t).strip()
            if not t:
                continue

            t = replace_doc_id_in_text(t)
            processed_texts.append(t)

        merged_text = "；".join(processed_texts)
        parts.append(merged_text)

    # =========================
    # 5️⃣ 输出
    # =========================
    final_text = "\n".join(parts)

    return final_text

# =========================
# 字体
# =========================
def set_dual_font_10pt(run):
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')


# =========================
# 替换
# =========================
def replace_text_strict_format(paragraph, replacements):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for k, v in replacements.items():
        new_text = new_text.replace(k, str(v))

    if new_text != full_text:
        paragraph.runs[0].text = new_text
        for i in range(1, len(paragraph.runs)):
            paragraph.runs[i].text = ""


def replace_all_content_keep_style(doc, replacements):
    for p in doc.paragraphs:
        replace_text_strict_format(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_strict_format(p, replacements)

    for node in doc.element.xpath("//w:t"):
        if node.text:
            text = node.text
            for k, v in replacements.items():
                text = text.replace(k, str(v))
            node.text = text
            
def replace_db_count(paragraph, placeholder, value):
    """
    将：检索出{万方检索量}篇文献
    替换为：检索出9篇文献
    且数字不加粗
    """

    if placeholder not in paragraph.text:
        return

    for run in paragraph.runs:
        if placeholder in run.text:
            parts = run.text.split(placeholder)

            # 清空当前 run
            run.text = ""

            # 前半段（继承原格式）
            run.add_text(parts[0])

            # 数字（强制不加粗）
            r = run._element
            new_run = paragraph.add_run(str(value))
            new_run.bold = False

            # 后半段
            if len(parts) > 1:
                paragraph.add_run(parts[1])



# =========================
# 页眉
# =========================
def replace_header_keep_format(doc, replacements):
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    for k, v in replacements.items():
                        run.text = run.text.replace(k, str(v))


# =========================
# ⭐ 表格填充（支持左对齐）
# =========================
def fill_specific_table(table, records):
    if not records or len(table.rows) < 2:
        return

    template_row = table.rows[1]

    while len(table.rows) > 2:
        table._tbl.remove(table.rows[-1]._tr)

    for idx, record in enumerate(records, start=1):
        new_row_xml = copy.deepcopy(template_row._tr)
        table._tbl.append(new_row_xml)
        row = table.rows[-1]

        row_data = [str(idx), record[1], record[2], record[3], record[4]]

        for i, val in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.text = ""

            run = p.add_run(str(val))
            set_dual_font_10pt(run)

            # ⭐ 第二列左对齐
            if i == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table._tbl.remove(template_row._tr)


# =========================
# 主程序
# =========================
from io import BytesIO

def generate_report(file_map, template_file):

    print("读取数据...")

    data = load_all_tables(file_map)
    
    replacements = load_placeholders_from_product_table(data["product_info"])

    databases = load_literature_from_screening(data["screening"])

    wanfang = databases.get("万方", [])
    cnki = databases.get("中国知网", [])
    pubmed = databases.get("PubMed", [])
    cochrane = databases.get("Cochrane Library", [])

    product_name = replacements.get("{产品名称}", "")

    total_included, summary_text = build_summary_text(databases, product_name)

    total_all = len(wanfang) + len(cnki) + len(pubmed) + len(cochrane)

    replacements["{文献量}"] = str(total_all)
    replacements["{总纳入文献量}"] = str(total_included)
    replacements["{文献数据分析}"] = summary_text

    replacements["{万方检索量}"] = len(wanfang)
    replacements["{知网检索量}"] = len(cnki)
    replacements["{PubMed检索量}"] = len(pubmed)
    replacements["{Cochrane检索量}"] = len(cochrane)

    analysis_text = build_literature_analysis_text(
        data.get("analysis"),
        data.get("screening")
    )

    replacements["{文献综述分析}"] = analysis_text

    doc = Document(BytesIO(template_file.read()))

    replace_all_content_keep_style(doc, replacements)
    replace_header_keep_format(doc, replacements)

    tables = doc.tables

    db_map = {
        3: wanfang,
        4: cnki,
        5: pubmed,
        6: cochrane
    }

    for idx, records in db_map.items():
        if len(tables) > idx:
            fill_specific_table(tables[idx], records)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output
